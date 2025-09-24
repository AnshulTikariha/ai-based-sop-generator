import re
import ast
import json
from typing import Dict, Any, List, Optional
import tempfile
import os


_CURL_SPLIT_RE = re.compile(r"\n\s*\n+", re.MULTILINE)


def _normalize_curls(curls_text: Optional[str], curls_list: Optional[List[str]]) -> List[str]:
    items: List[str] = []
    if curls_list:
        items.extend([c for c in curls_list if isinstance(c, str) and c.strip()])
    if curls_text and curls_text.strip():
        # split on blank lines if multiple curls in one text block
        chunks = _CURL_SPLIT_RE.split(curls_text.strip())
        for ch in chunks:
            if ch.strip().startswith("curl "):
                items.append(ch.strip())
            else:
                # Keep non-curl chunks for potential JSON fallback
                items.append(ch.strip())
    # de-dup and keep order
    seen = set()
    unique: List[str] = []
    for c in items:
        if c not in seen:
            unique.append(c)
            seen.add(c)
    return unique


def _parse_headers(tokens: List[str]) -> Dict[str, str]:
    headers: Dict[str, str] = {}
    for i, tok in enumerate(tokens):
        if tok in ("-H", "--header") and i + 1 < len(tokens):
            raw = tokens[i + 1].strip().strip("'\"")
            if ":" in raw:
                k, v = raw.split(":", 1)
                headers[k.strip()] = v.strip()
    return headers


def _parse_data(tokens: List[str]) -> Optional[str]:
    for i, tok in enumerate(tokens):
        if tok in ("-d", "--data", "--data-raw", "--data-binary") and i + 1 < len(tokens):
            raw = tokens[i + 1].strip().strip("'\"")
            # Attempt to unescape common shell escaping of JSON for better parsing later
            try:
                # Replace backslash-escaped quotes if present
                cleaned = raw.encode('utf-8').decode('unicode_escape')
            except Exception:
                cleaned = raw
            return cleaned
    return None


def _coerce_json(value: str) -> Any:
    """Best-effort: turn variously escaped JSON-ish strings into Python objects.

    Tries json.loads; if that fails, progressively normalizes common cURL/shell escaping patterns
    and tries again. Falls back to ast.literal_eval for edge cases.
    """
    if not isinstance(value, str):
        return value
    txt = value.strip()
    # Fast path
    try:
        return json.loads(txt)
    except Exception:
        pass
    # Common normalizations
    candidates = []
    # 0) Extract JSON substring between first '{' and last '}' if present
    if '{' in txt and '}' in txt:
        start = txt.find('{')
        end = txt.rfind('}') + 1
        if end > start:
            candidates.append(txt[start:end])
    candidates.append(txt.replace('\\"', '"'))
    candidates.append(re.sub(r"\\/", "/", txt))
    candidates.append(re.sub(r"\\n", "", txt))
    # Remove stray backslashes before quotes/braces
    candidates.append(re.sub(r'\\([{}\[\]"])', r"\1", txt))
    for cand in candidates:
        try:
            return json.loads(cand)
        except Exception:
            continue
    # Try literal_eval for python-like dicts
    try:
        obj = ast.literal_eval(txt)
        return obj
    except Exception:
        return value


def _parse_method(tokens: List[str]) -> Optional[str]:
    for i, tok in enumerate(tokens):
        if tok in ("-X", "--request") and i + 1 < len(tokens):
            return tokens[i + 1].strip().upper()
    return None


def _parse_url(tokens: List[str]) -> Optional[str]:
    # URL is usually the last non-flag token or after --url
    url: Optional[str] = None
    for i, tok in enumerate(tokens):
        if tok == "--url" and i + 1 < len(tokens):
            url = tokens[i + 1]
    if not url:
        # pick last token that looks like a URL
        for tok in reversed(tokens):
            if tok.startswith("http://") or tok.startswith("https://") or tok.startswith("/"):
                url = tok
                break
    if url:
        return url.strip().strip("'\"")
    return None


def _shell_split(s: str) -> List[str]:
    # lightweight splitter respecting single/double quotes
    parts: List[str] = []
    buf: List[str] = []
    quote: Optional[str] = None
    i = 0
    while i < len(s):
        ch = s[i]
        if quote:
            if ch == quote:
                quote = None
            else:
                buf.append(ch)
        else:
            if ch in ("'", '"'):
                quote = ch
            elif ch.isspace():
                if buf:
                    parts.append(''.join(buf))
                    buf = []
            else:
                buf.append(ch)
        i += 1
    if buf:
        parts.append(''.join(buf))
    return parts


def parse_curl(curl_cmd: str) -> Dict[str, Any]:
    tokens = _shell_split(curl_cmd.strip().lstrip("curl "))
    body_peek = _parse_data(tokens)
    method = _parse_method(tokens) or ("POST" if body_peek is not None else "GET")
    url = _parse_url(tokens) or "/"
    headers = _parse_headers(tokens)
    data = body_peek
    return {"method": method, "url": url, "headers": headers, "body": data}


def parse_curl_inputs(payload: Any) -> Dict[str, Any]:
    curls = _normalize_curls(getattr(payload, "curls_text", None), getattr(payload, "curls", None))
    requests: List[Dict[str, Any]] = []
    leftovers: List[str] = []
    for c in curls:
        if c.startswith("curl "):
            requests.append(parse_curl(c))
        else:
            leftovers.append(c)
    # Fallback: if there are no valid curl requests but the text looks like JSON, create a default POST
    if not requests and leftovers:
        combined = "\n".join(leftovers).strip()
        # Try to find a JSON object or array
        json_str = combined
        try:
            parsed_json = json.loads(json_str)
            requests.append({
                "method": "POST",
                "url": "/",
                "headers": {"Content-Type": "application/json"},
                "body": json.dumps(parsed_json)
            })
        except Exception:
            # Ignore if not JSON
            pass
    return {"requests": requests}


def _infer_base_url(requests: List[Dict[str, Any]], hint: Optional[str]) -> Optional[str]:
    if hint:
        return hint.rstrip('/')
    for r in requests:
        url = r.get("url") or ""
        m = re.match(r"^(https?://[^/]+)", url)
        if m:
            return m.group(1).rstrip('/')
    return None


def _path_from_url(url: str, base_url: Optional[str]) -> str:
    if base_url and url.startswith(base_url):
        return url[len(base_url):] or "/"
    # strip scheme+host if present
    m = re.match(r"^https?://[^/]+(.*)$", url)
    if m:
        return m.group(1) or "/"
    return url if url.startswith('/') else "/" + url


def _mask_sensitive(headers: Dict[str, str]) -> Dict[str, str]:
    masked: Dict[str, str] = {}
    for k, v in headers.items():
        if k.lower() in ("authorization", "x-api-key", "api-key"):
            masked[k] = "***"
        else:
            masked[k] = v
    return masked


def build_openapi_from_requests(project_name: str, base_url_hint: Optional[str], requests_payload: List[Dict[str, Any]], ai_enabled: bool = True) -> Dict[str, Any]:
    base_url = _infer_base_url(requests_payload, base_url_hint)
    paths: Dict[str, Any] = {}
    uses_bearer = False
    uses_api_key = False

    for r in requests_payload:
        method = r.get("method", "GET").lower()
        url = r.get("url", "/")
        path = _path_from_url(url, base_url)
        headers = _mask_sensitive(r.get("headers", {}))
        body = r.get("body")

        # Propose a tag from first non-empty path segment
        tag = None
        segs = [s for s in path.split('/') if s]
        if segs:
            tag = segs[0]

        op: Dict[str, Any] = {
            "summary": f"{method.upper()} {path}",
            "operationId": re.sub(r"[^a-zA-Z0-9]", "_", f"{method}_{path}").strip("_"),
            "tags": [tag or "general"],
            "parameters": [],
            "responses": {
                "200": {
                    "description": "OK",
                    "content": {"application/json": {"schema": {"type": "object"}}}
                }
            }
        }

        # query parameters
        if "?" in path:
            path_only, query = path.split("?", 1)
            path = path_only
            for q in query.split("&"):
                if not q:
                    continue
                name = q.split("=", 1)[0]
                op["parameters"].append({
                    "in": "query",
                    "name": name,
                    "schema": {"type": "string"}
                })

        # path parameters heuristic
        for seg in path.split('/'):
            if re.fullmatch(r"[0-9a-fA-F-]{6,}", seg) or seg.isdigit():
                # turn into templated path param
                path = path.replace("/" + seg, "/{id}")
                op["parameters"].append({
                    "in": "path",
                    "name": "id",
                    "required": True,
                    "schema": {"type": "string"}
                })
                break

        # headers
        for hk, hv in headers.items():
            if hk.lower() == "authorization" and isinstance(hv, str) and hv.lower().startswith("bearer"):
                uses_bearer = True
                continue
            if hk.lower() in ("x-api-key", "api-key"):
                uses_api_key = True
                continue
            op["parameters"].append({
                "in": "header",
                "name": hk,
                "schema": {"type": "string"},
                "example": hv
            })

        # requestBody
        if body:
            example = None
            schema: Dict[str, Any] = {"type": "string"}
            # Try robust JSON coercion
            coerced = _coerce_json(body) if isinstance(body, str) else body
            if isinstance(coerced, (dict, list)):
                example = coerced
                schema = {"type": "object" if isinstance(coerced, dict) else "array"}
            else:
                try:
                    example = json.loads(body)
                    schema = {"type": "object"}
                except Exception:
                    example = body
            op["requestBody"] = {
                "content": {
                    "application/json": {
                        "schema": schema,
                        "example": example
                    }
                }
            }

        # Standard error responses if not already present
        for code, desc in {
            "400": "Bad Request",
            "401": "Unauthorized",
            "403": "Forbidden",
            "404": "Not Found",
            "429": "Too Many Requests",
            "500": "Internal Server Error",
        }.items():
            if code not in op["responses"]:
                op["responses"][code] = {"description": desc}

        paths.setdefault(path, {})[method] = op

    components: Dict[str, Any] = {"schemas": {}}
    security: List[Dict[str, Any]] = []
    if uses_bearer:
        components.setdefault("securitySchemes", {})["bearerAuth"] = {
            "type": "http",
            "scheme": "bearer",
            "bearerFormat": "JWT"
        }
        security.append({"bearerAuth": []})
    if uses_api_key:
        components.setdefault("securitySchemes", {})["apiKeyAuth"] = {
            "type": "apiKey",
            "in": "header",
            "name": "X-API-Key"
        }
        security.append({"apiKeyAuth": []})

    openapi: Dict[str, Any] = {
        "openapi": "3.0.3",
        "info": {"title": f"{project_name} API", "version": "0.1.0"},
        "servers": ([{"url": base_url}] if base_url else []),
        "paths": paths,
        "components": components,
        **({"security": security} if security else {})
    }

    if ai_enabled:
        # Until a local HF model is fully available, synthesize a precise description deterministically
        for p, ops in paths.items():
            for m, op in ops.items():
                header_names = [pr.get('name') for pr in op.get('parameters',[]) if pr.get('in')=='header']
                query_names = [pr.get('name') for pr in op.get('parameters',[]) if pr.get('in')=='query']
                body_example = op.get('requestBody',{}).get('content',{}).get('application/json',{}).get('example')
                body_keys = list(body_example.keys())[:12] if isinstance(body_example, dict) else []
                # Build concise sentence
                parts: List[str] = []
                verb = 'Retrieves' if m.upper() == 'GET' else 'Processes'
                parts.append(f"{verb} data for `{p}`")
                if query_names:
                    parts.append(f"with query filters {query_names}")
                if body_keys:
                    parts.append(f"using body fields {body_keys}")
                sentence = ' '.join(parts) + "."
                op["description"] = sentence[:800]

    return openapi


def _render_sheet_style(openapi: Dict[str, Any]) -> str:
    lines: List[str] = []
    info = openapi.get("info", {})
    title = info.get("title", "API")
    base = openapi.get("servers", [{}])
    base_url = base[0].get("url") if base else None
    lines.append(f"# {title}")
    lines.append("")
    if base_url:
        lines.append(f"**Base URL**: `{base_url}`")
    # Authentication overview
    sec = openapi.get("security") or []
    if sec:
        lines.append("")
        lines.append("**Authentication**:")
        if any("bearerAuth" in s for s in sec):
            lines.append("- Bearer token (JWT) via `Authorization: Bearer <token>` header")
        if any("apiKeyAuth" in s for s in sec):
            lines.append("- API Key via `X-API-Key: <key>` header")
    lines.append("")
    lines.append("> This document is generated automatically from provided cURL requests.")
    lines.append("")
    for path, ops in openapi.get("paths", {}).items():
        for method, op in ops.items():
            lines.append(f"## {op.get('summary') or method.upper() + ' ' + path}")
            lines.append("")
            lines.append(f"- Method: **{method.upper()}**")
            lines.append(f"- URL: `{path}`")
            if op.get("tags"):
                lines.append(f"- Tags: {', '.join(op['tags'])}")
            if op.get("parameters"):
                headers = [p for p in op["parameters"] if p.get("in") == "header"]
                query = [p for p in op["parameters"] if p.get("in") == "query"]
                path_params = [p for p in op["parameters"] if p.get("in") == "path"]
                if headers:
                    lines.append("\n### Headers")
                    lines.append("| Name | Example |\n|---|---|")
                    for h in headers:
                        lines.append(f"| {h.get('name')} | {str(h.get('example') or '')} |")
                if path_params:
                    lines.append("\n### Path Params")
                    lines.append("| Name | Required |\n|---|---|")
                    for p in path_params:
                        lines.append(f"| {p.get('name')} | { 'yes' if p.get('required') else 'no' } |")
                if query:
                    lines.append("\n### Query Params")
                    lines.append("| Name |\n|---|")
                    for q in query:
                        lines.append(f"| {q.get('name')} |")
            if op.get("requestBody"):
                lines.append("\n### Request Body")
                example = op["requestBody"].get("content", {}).get("application/json", {}).get("example")
                # If example is a string that looks like JSON, parse/normalize it first
                if isinstance(example, str):
                    example = _coerce_json(example)
                if example is not None:
                    pretty = json.dumps(example, indent=2) if not isinstance(example, str) else example
                    lines.append("""
```json
%s
```
""".strip() % pretty)
                # Field dictionary (professional payload description)
                def _infer_type(value: Any) -> str:
                    if value is None:
                        return "String"
                    if isinstance(value, bool):
                        return "Boolean"
                    if isinstance(value, int) or isinstance(value, float):
                        return "Number"
                    if isinstance(value, str):
                        return "String"
                    if isinstance(value, list):
                        inner = _infer_type(value[0]) if value else "Any"
                        return f"Array<{inner}>"
                    if isinstance(value, dict):
                        return "Object"
                    return "String"

                def _infer_desc(key: str, value: Any) -> str:
                    k = key.lower()
                    # Domain-specific heuristics
                    domain_map = {
                        "company_id": "Company identifier.",
                        "user_id": "User identifier initiating the request.",
                        "module_id": "Module identifier.",
                        "user_role_id": "Role identifier for authorization decisions.",
                        "module_name": "Human-readable module name.",
                        "page_no": "Page number for pagination (1-based).",
                        "record_per_page": "Number of records per page.",
                        "force_active_status": "If '1', restrict results to active entities.",
                        "employee_ids": "List of employee identifiers to filter results.",
                        "url": "Client page/route where the request originated.",
                        "device_info": "Information about client device and environment.",
                        "device_type": "Type of device (Desktop/Mobile/Tablet).",
                        "is_mobile": "Flag indicating mobile device (0/1).",
                        "is_tablet": "Flag indicating tablet device (0/1).",
                        "is_desktop": "Flag indicating desktop device (0/1).",
                        "browser": "Client browser name.",
                        "os": "Operating system name.",
                        "os_version": "Operating system version.",
                        "user_agent": "Raw user-agent string.",
                        "ip_address": "Client IP address.",
                        "is_vendor": "Flag indicating vendor context (0/1).",
                    }
                    if k in domain_map:
                        return domain_map[k]
                    if "email" in k:
                        return "User email address."
                    if k.endswith("_id") or k == "id":
                        return "Unique identifier."
                    if "phone" in k or "mobile" in k:
                        return "Phone number."
                    if "name" in k:
                        return "Descriptive name."
                    if "password" in k or "passcode" in k:
                        return "Secret credential; do not log."
                    if "gst" in k or "gstin" in k:
                        return "GST identification number."
                    if "pan" in k:
                        return "PAN number."
                    if "account" in k and "number" in k:
                        return "Bank account number (mask for security)."
                    if "currency" in k:
                        return "Currency code (e.g., INR, USD)."
                    if isinstance(value, list):
                        return "List of values."
                    if isinstance(value, dict):
                        return "Nested object."
                    # Fallback generic
                    return f"Field '{key}'."

                def _flatten(prefix: str, obj: Any, out: List[Dict[str, str]]):
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            _flatten(f"{prefix}.{k}" if prefix else k, v, out)
                    elif isinstance(obj, list):
                        sample = obj[0] if obj else None
                        t = _infer_type(obj)
                        out.append({
                            "name": prefix,
                            "type": t,
                            "description": _infer_desc(prefix.split('.')[-1], obj)
                        })
                        if isinstance(sample, (dict, list)):
                            _flatten(prefix + "[]", sample, out)
                    else:
                        out.append({
                            "name": prefix,
                            "type": _infer_type(obj),
                            "description": _infer_desc(prefix.split('.')[-1], obj)
                        })

                if isinstance(example, (dict, list)):
                    fields: List[Dict[str, str]] = []
                    _flatten("", example, fields)
                    if fields:
                        lines.append("\n### Request Body Fields")
                        lines.append("| Field | Type | Description |\n|---|---|---|")
                        for f in fields:
                            desc = f.get("description") or ""
                            lines.append(f"| `{f['name']}` | {f['type']} | {desc} |")
            # Examples
            # Build curl from info we have
            curl_parts = ["curl", "-X", method.upper()]
            full_url = (base_url or "") + path
            curl_parts.append(f"\"{full_url}\"")
            lines.append("\n### Example cURL")
            lines.append("""
```bash
%s
```
""".strip() % (" ".join(curl_parts)))
            if op.get("responses"):
                lines.append("\n### Responses")
                lines.append("| Status | Description |\n|---|---|")
                for code, resp in op["responses"].items():
                    lines.append(f"| {code} | {resp.get('description','')} |")
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_markdown_from_openapi(openapi: Dict[str, Any], style: str = "default") -> str:
    if style == "sheet":
        return _render_sheet_style(openapi)
    if style == "vendor":
        return _render_vendor_style(openapi)
    lines: List[str] = []
    info = openapi.get("info", {})
    title = info.get("title", "API")
    lines.append(f"# {title}")
    lines.append("")
    if openapi.get("servers"):
        lines.append("Servers:")
        for s in openapi["servers"]:
            lines.append(f"- {s.get('url')}")
        lines.append("")

    for path, ops in openapi.get("paths", {}).items():
        lines.append(f"## {path}")
        for method, op in ops.items():
            lines.append(f"### {method.upper()}")
            if op.get("summary"):
                lines.append(op["summary"])
            if op.get("parameters"):
                lines.append("")
                lines.append("Parameters:")
                for p in op["parameters"]:
                    loc = p.get("in")
                    name = p.get("name")
                    lines.append(f"- {loc} `{name}`")
            if op.get("requestBody"):
                lines.append("")
                lines.append("Request Body:")
                example = op["requestBody"].get("content", {}).get("application/json", {}).get("example")
                if example is not None:
                    pretty = json.dumps(example, indent=2) if not isinstance(example, str) else example
                    lines.append("""
```json
%s
```
""".strip() % pretty)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _render_vendor_style(openapi: Dict[str, Any]) -> str:
    lines: List[str] = []
    info = openapi.get("info", {})
    title = info.get("title", "API")
    # Avoid duplicated 'API' in the header (e.g., 'Generated API API')
    try:
        display_title = re.sub(r"\s+API$", "", title)
    except Exception:
        display_title = title
    version = info.get("version", "1.0")
    base = openapi.get("servers", [{}])
    base_url = base[0].get("url") if base else None

    # Header
    lines.append(f"# API Documentation: {display_title}")
    lines.append("")
    lines.append(f"Version: `{version}`")
    if base_url:
        lines.append(f"Base URL: `{base_url}`")
    lines.append("")

    # Auth section
    sec = openapi.get("security") or []
    if sec:
        lines.append("## Authentication")
        if any("bearerAuth" in s for s in sec):
            lines.append("- Bearer token via `Authorization: Bearer <token>` header")
        if any("apiKeyAuth" in s for s in sec):
            lines.append("- API Key via `X-API-Key: <key>` header")
        lines.append("")

    # Standards
    lines.append("## Conventions")
    lines.append("- Content-Type: application/json")
    lines.append("- Date/time in ISO 8601, UTC")
    lines.append("- Idempotency: GET safe; POST/PUT/PATCH/DELETE may change state")
    lines.append("")

    # Paths
    # We intentionally avoid static descriptions; prefer AI/populated OpenAPI fields

    for path, ops in openapi.get("paths", {}).items():
        for method, op in ops.items():
            # Title per endpoint
            if 'ajax_getempcostcenter' in path:
                lines.append("## Get Employee Cost Center")
            else:
                lines.append(f"## {op.get('summary') or method.upper() + ' ' + path}")
            lines.append("")
            # Endpoint block
            lines.append("### Endpoint")
            lines.append(f"`{method.upper()} {path}`")
            if op.get("tags"):
                lines.append(f"- Tags: {', '.join(op['tags'])}")
            # Description: use OpenAPI op.description if present; otherwise omit
            desc = (op.get("description") or "").strip()
            lines.append("")
            lines.append("### Description")
            # Ensure description is a paragraph and left-aligned; fall back text if empty
            safe_desc = desc if desc else "(description not provided)"
            lines.append(safe_desc)

            # Parameters
            params = op.get("parameters", [])
            headers = [p for p in params if p.get("in") == "header"]
            path_params = [p for p in params if p.get("in") == "path"]
            query_params = [p for p in params if p.get("in") == "query"]
            if path_params:
                lines.append("\n### Path Parameters")
                lines.append("| Name | Required | Description |\n|---|---|---|")
                for p in path_params:
                    lines.append(f"| `{p.get('name')}` | {'Yes' if p.get('required') else 'No'} | |")
            if query_params:
                lines.append("\n### Query Parameters")
                lines.append("| Name | Type | Required | Description |\n|---|---|---|---|")
                for p in query_params:
                    typ = (p.get('schema') or {}).get('type', 'string')
                    lines.append(f"| `{p.get('name')}` | {typ} | No | |")
            # Headers: derive only from provided cURL/OpenAPI parameters
            params = op.get("parameters", [])
            headers = [p for p in params if p.get("in") == "header"]
            if headers:
                lines.append("\n### Headers")
                # Per request: remove Description column from headers table
                lines.append("| Header | Type | Required |\n|---|---|---|")
                for h in headers:
                    nm = h.get('name')
                    req = 'Yes' if h.get('required') else 'No'
                    lines.append(f"| {nm} | String | {req} |")

            # Request body
            rb = op.get("requestBody")
            if rb:
                lines.append("\n### Request Body")
                example = rb.get("content", {}).get("application/json", {}).get("example")
                if isinstance(example, str):
                    example = _coerce_json(example)
                if example is not None:
                    pretty = json.dumps(example, indent=2) if not isinstance(example, str) else example
                    lines.append("""
```json
%s
```
""".strip() % pretty)

                # Field table
                def _infer_type(value: Any) -> str:
                    if value is None:
                        return "String"
                    if isinstance(value, bool):
                        return "Boolean"
                    if isinstance(value, int) or isinstance(value, float):
                        return "Number"
                    if isinstance(value, str):
                        return "String"
                    if isinstance(value, list):
                        inner = _infer_type(value[0]) if value else "Any"
                        return f"Array<{inner}>"
                    if isinstance(value, dict):
                        return "Object"
                    return "String"

                def _infer_desc(key: str, value: Any) -> str:
                    k = key.lower()
                    # Extended, domain-aware descriptions
                    domain_map = {
                        "company_id": "ID of the company.",
                        "user_id": "User ID making the request.",
                        "user_role_id": "Role ID of the logged-in user.",
                        "page_no": "Page number for pagination.",
                        "record_per_page": "Number of records per page.",
                        "force_active_status": "Filter for active status (1 = active only).",
                        "employee_ids": "List of employee IDs to filter.",
                        "url": "Module/section reference URL.",
                        "is_vendor": "0 = employee, 1 = vendor.",
                        "module_name": "Name of the module accessing API.",
                        "module_id": "ID of the module.",
                        "pr_no": "Purchase Request numbers filter.",
                        "po_no": "Purchase Order numbers filter.",
                        "irn_no": "IRN numbers filter.",
                        "vendor_id": "Vendor identifiers filter.",
                        "vendor_invoice_no": "Vendor invoice number filter.",
                        "vendor_invoice_date": "Vendor invoice date filter (YYYY-MM-DD).",
                        "invoice_approved_by": "Approver user IDs filter.",
                        "invoice_approved_date": "Invoice approval date filter.",
                        "gross_invoice_amount": "Gross invoice amount filter.",
                        "tds_amount": "TDS amount filter.",
                        "advance_deducted": "Advance deducted amount filter.",
                        "net_payable_amount": "Net payable amount filter.",
                        "payment_requested_by": "Payment requesting user IDs filter.",
                        "payment_entry_no": "Payment entry number filter.",
                        "payment_entry_date": "Payment entry date filter.",
                        "payment_amount": "Payment amount filter.",
                        "remaining_payment_amount": "Remaining payment amount filter.",
                        "payment_mode": "Payment mode (e.g., NEFT/RTGS/Cheque).",
                        "instrument_no": "Instrument/cheque number.",
                        "company_bank": "Company bank name.",
                        "utr_no": "Bank UTR number.",
                        "from_amount_clearing_date": "Start date for amount clearing range.",
                        "to_amount_clearing_date": "End date for amount clearing range.",
                        "payment_approval_status": "Payment approval status filter.",
                        "payment_status": "Payment status filter.",
                        "grn_approval_no": "GRN approval number.",
                        "column_list": "Columns to include in the report output.",
                        "from_payment_date": "Start payment date range.",
                        "to_payment_date": "End payment date range.",
                        "search_query": "Free-text search query.",
                        "is_excel_download": "If '1', export as Excel instead of JSON.",
                        "invoice_payment_remarks": "Remarks filter for invoice payments.",
                        "request_server_time": "Client-side request time (for logging).",
                        "device_info": "Device/browser metadata.",
                        "device_type": "Type of device (Desktop, Mobile, etc.).",
                        "is_mobile": "1 = mobile, 0 = not mobile.",
                        "is_tablet": "1 = tablet, 0 = not tablet.",
                        "is_desktop": "1 = desktop, 0 = not desktop.",
                        "browser": "Browser name (e.g., Chrome).",
                        "os": "Operating system (e.g., Windows/Mac).",
                        "os_version": "OS version (e.g., windows-10).",
                        "user_agent": "Full user agent string.",
                        "ip_address": "Client IP address (if available).",
                    }
                    if k in domain_map:
                        return domain_map[k]
                    if "email" in k:
                        return "Email address."
                    if k.endswith("_id") or k == "id":
                        return "Unique identifier."
                    if "phone" in k or "mobile" in k:
                        return "Phone number."
                    if "name" in k:
                        return "Descriptive name."
                    if "password" in k:
                        return "Secret credential; never log."
                    if "gst" in k:
                        return "GST identification number."
                    if "pan" in k:
                        return "PAN number."
                    if "account" in k and "number" in k:
                        return "Bank account number (mask in logs)."
                    if "currency" in k:
                        return "Currency code (e.g., INR, USD)."
                    if isinstance(value, list):
                        return "List of values."
                    if isinstance(value, dict):
                        return "Nested object."
                    return ""

                def _flatten(prefix: str, obj: Any, out: List[Dict[str, str]]):
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            _flatten(f"{prefix}.{k}" if prefix else k, v, out)
                    elif isinstance(obj, list):
                        sample = obj[0] if obj else None
                        out.append({
                            "name": prefix,
                            "type": _infer_type(obj),
                            "description": _infer_desc(prefix.split('.')[-1], obj)
                        })
                        if isinstance(sample, (dict, list)):
                            _flatten(prefix + "[]", sample, out)
                    else:
                        out.append({
                            "name": prefix,
                            "type": _infer_type(obj),
                            "description": _infer_desc(prefix.split('.')[-1], obj)
                        })

                if isinstance(example, (dict, list)):
                    fields: List[Dict[str, str]] = []
                    _flatten("", example, fields)
                    if fields:
                        lines.append("\n### Request Body Fields")
                        lines.append("| Field | Type | Required | Description |\n|---|---|---|---|")
                        required_keys = {"company_id","user_id","user_role_id"}
                        for f in fields:
                            req = "Yes" if f['name'].split('.')[-1] in required_keys else "No"
                            lines.append(f"| `{f['name']}` | {f['type']} | {req} | {f.get('description','')} |")

                        # Nested device_info table if present
                        if isinstance(example, dict) and isinstance(example.get('device_info'), dict):
                            dev = example['device_info']
                            lines.append("\n#### device_info object")
                            lines.append("| Field | Type | Description |\n|---|---|---|")
                            for k, v in dev.items():
                                lines.append(f"| `{k}` | {_infer_type(v)} | {_infer_desc(k, v)} |")

            # Responses
            if op.get("responses"):
                lines.append("\n### Responses")
                lines.append("| Status | Description |\n|---|---|")
                for code, resp in (op.get("responses") or {}).items():
                    lines.append(f"| {code} | {resp.get('description','')} |")

            # We intentionally do not include a sample request block
            lines.append("")

    # Support section
    lines.append("## Support & SLA")
    lines.append("- Response time targets: 99.9% uptime, <300ms P50 for core endpoints")
    lines.append("- Contact: support@example.com")
    lines.append("- Changelog: maintained by provider")
    lines.append("")

    return "\n".join(lines).strip() + "\n"


def generate_pdf(markdown_text: str) -> str:
    """Generate a professional-looking PDF from markdown using reportlab platypus.

    - Headings mapped to larger fonts
    - Paragraph spacing
    - Markdown tables rendered as bordered tables
    - Code blocks rendered in monospaced boxes
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import Preformatted

    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)

    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='H1', parent=styles['Heading1'], fontSize=24, leading=28, spaceAfter=14, textColor=colors.HexColor('#0f172a')))
    styles.add(ParagraphStyle(name='H2', parent=styles['Heading2'], fontSize=18, leading=22, spaceAfter=12, textColor=colors.HexColor('#111827')))
    styles.add(ParagraphStyle(name='H3', parent=styles['Heading3'], fontSize=14, leading=20, spaceAfter=10, textColor=colors.HexColor('#111827')))
    styles.add(ParagraphStyle(name='Body', parent=styles['BodyText'], fontSize=12, leading=18, spaceAfter=8, textColor=colors.black))
    # Use a unique style name to avoid conflicts with default styles
    # Use light background and dark text to match request for black text
    styles.add(ParagraphStyle(name='CodeBlock', fontName='Courier', fontSize=11, leading=16, backColor=colors.HexColor('#f3f4f6'), textColor=colors.black, leftIndent=8, rightIndent=8, spaceAfter=10, borderPadding=8))

    lines = markdown_text.split('\n')
    elements: List[Any] = []

    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue

        # Code block ```
        if line.strip().startswith('```'):
            fence = line.strip()
            code_lines: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # skip closing fence
            i += 1
            elements.append(Preformatted('\n'.join(code_lines), styles['CodeBlock']))
            continue

        # Table block (markdown table starts with |... and has a separator on next line)
        if '|' in line and line.strip().startswith('|'):
            # collect contiguous table lines
            tbl_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            # Parse table
            rows = []
            for tl in tbl_lines:
                # remove leading/trailing |
                row = [c.strip() for c in tl.strip().strip('|').split('|')]
                # skip separator rows of ---
                if all(set(c) <= set('-: ') and c for c in row):
                    continue
                rows.append(row)
            if rows:
                # Make table span full available width
                num_cols = len(rows[0]) if rows else 1
                # Build Paragraph-wrapped cells to allow multiline wrapping
                header_style = ParagraphStyle(name='TblHead', parent=styles['Body'], fontName='Helvetica-Bold', textColor=colors.white)
                cell_style = ParagraphStyle(name='TblCell', parent=styles['Body'], fontSize=11, leading=14)
                wrapped_rows: List[List[Any]] = []
                for ridx, r in enumerate(rows):
                    new_row: List[Any] = []
                    for c in r:
                        txt = re.sub(r'\s+', ' ', c)
                        if ridx == 0:
                            new_row.append(Paragraph(txt, header_style))
                        else:
                            new_row.append(Paragraph(txt, cell_style))
                    wrapped_rows.append(new_row)
                # Prefer wider last column when header contains Description
                header_texts = [str(x) for x in rows[0]]
                clean_headers = [re.sub(r'<.*?>','',t).strip().lower() for t in header_texts]
                if num_cols == 4 and clean_headers == ['field','type','required','description']:
                    # Fix field/type columns wider so they don't shrink
                    col_widths = [doc.width*0.28, doc.width*0.18, doc.width*0.12, doc.width*0.42]
                elif any('description' in h.lower() for h in header_texts) and num_cols >= 4:
                    col_widths = [doc.width*0.25, doc.width*0.18, doc.width*0.12] + [doc.width*0.45]
                else:
                    col_width = doc.width / max(1, num_cols)
                    col_widths = [col_width]*num_cols
                # Use repeatRows=1 for header; allow word wrapping
                t = Table(wrapped_rows, colWidths=col_widths, hAlign='LEFT', repeatRows=1)
                t.setStyle(TableStyle([
                    ('FONT', (0,0), (-1,0), 'Helvetica-Bold', 12),
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1f2937')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('FONT', (0,1), (-1,-1), 'Helvetica', 11),
                    ('GRID', (0,0), (-1,-1), 0.6, colors.HexColor('#9ca3af')),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 8),
                    ('RIGHTPADDING', (0,0), (-1,-1), 8),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f3f4f6')]),
                    ('WORDWRAP', (0,0), (-1,-1), True),
                ]))
                elements.append(t)
                elements.append(Spacer(1, 12))
            continue

        # Headings
        if line.startswith('# '):
            elements.append(Paragraph(line[2:].strip(), styles['H1']))
        elif line.startswith('## '):
            elements.append(Paragraph(line[3:].strip(), styles['H2']))
        elif line.startswith('### '):
            elements.append(Paragraph(line[4:].strip(), styles['H3']))
        else:
            # Bullet lists
            if line.strip().startswith('- '):
                # group list items
                items = [line.strip()[2:]]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('- '):
                    items.append(lines[i].strip()[2:])
                    i += 1
                elements.append(Paragraph('• ' + '<br/>• '.join(items), styles['Body']))
                continue
            elements.append(Paragraph(line.strip(), styles['Body']))
        i += 1

    try:
        doc.build(elements)
        return path
    except Exception as e:
        # Fallback to basic PDF writer to avoid blocking downloads
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import simpleSplit
        fd2, path2 = tempfile.mkstemp(suffix=".pdf")
        os.close(fd2)
        c = canvas.Canvas(path2, pagesize=letter)
        width, height = letter
        margin = 40
        y = height - margin
        max_width = width - margin * 2
        for line in markdown_text.split('\n'):
            wrapped = simpleSplit(line, 'Helvetica', 10, max_width)
            for w in wrapped:
                if y < margin:
                    c.showPage()
                    y = height - margin
                c.setFont('Helvetica', 10)
                c.drawString(margin, y, w)
                y -= 14
        c.save()
        return path2


def generate_docx(markdown_text: str) -> str:
    """Generate a DOCX with headings and paragraphs from markdown.

    Uses python-docx; if not available, creates a plain-text DOCX-like via fallback.
    """
    try:
        from docx import Document  # type: ignore
        import docx
    except Exception as e:
        # Fallback: write .docx as plain text is not viable; require python-docx
        raise RuntimeError("python-docx not installed. Please add python-docx to requirements.")

    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.strip().startswith('```'):
            # collect code block
            i += 1
            code_lines: List[str] = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph('\n'.join(code_lines))
            p.runs[0].font.name = 'Courier New'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        # Markdown table detection
        if '|' in line and line.strip().startswith('|'):
            tbl_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            rows: List[List[str]] = []
            for tl in tbl_lines:
                row = [c.strip() for c in tl.strip().strip('|').split('|')]
                # skip markdown separator row
                if all(set(c) <= set('-: ') and c for c in row):
                    continue
                rows.append(row)
            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=0, cols=cols)
                table.autofit = False
                total_width = section.page_width - section.left_margin - section.right_margin
                col_w = int(total_width / max(1, cols))
                for irow, r in enumerate(rows):
                    tr = table.add_row().cells
                    for j, cell_text in enumerate(r):
                        tr[j].text = cell_text
                        tr[j].width = col_w
                # Header style
                if table.rows:
                    hdr = table.rows[0]
                    for cell in hdr.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
                # Set borders
                tbl = table._tbl
                tblPr = tbl.tblPr
                tblPr.append(docx.oxml.parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="6" w:space="0" w:color="9ca3af"/><w:left w:val="single" w:sz="6" w:space="0" w:color="9ca3af"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="9ca3af"/><w:right w:val="single" w:sz="6" w:space="0" w:color="9ca3af"/><w:insideH w:val="single" w:sz="6" w:space="0" w:color="9ca3af"/><w:insideV w:val="single" w:sz="6" w:space="0" w:color="9ca3af"/></w:tblBorders>'))
            continue
        # default paragraph
        doc.add_paragraph(line.strip())
        i += 1
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    os.remove(path)
    doc.save(path)
    return path


